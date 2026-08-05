import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Styles
title_style = doc.styles['Title']
title_style.font.name = 'Arial'
title_style.font.size = Pt(24)

heading1_style = doc.styles['Heading 1']
heading1_style.font.name = 'Arial'
heading1_style.font.size = Pt(16)

normal_style = doc.styles['Normal']
normal_style.font.name = 'Arial'
normal_style.font.size = Pt(11)

# Title Page
doc.add_heading('The Lighthouse', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('A Self-Reflection and Wellbeing Platform\n\nVisual Lab Report', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

def add_section(title, description, figure_caption):
    doc.add_heading(title, level=1)
    if description:
        doc.add_paragraph(description, style='Normal')
    
    # Add empty space for image
    p = doc.add_paragraph()
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    
    # Set row height to 3.5 inches
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '5040') # 3.5 inches in twips (1440 * 3.5)
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)
    
    # Add a prompt in the table
    cell = table.cell(0, 0)
    cell.text = "\n\n\n\n[PASTE SCREENSHOT HERE]"
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Caption
    cap = doc.add_paragraph(figure_caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].bold = True
    doc.add_paragraph() # spacing

# Introduction
doc.add_heading('1. Project Overview', level=1)
doc.add_paragraph(
    "The Lighthouse is a comprehensive wellbeing platform designed to help users track their emotional state, "
    "reflect on their daily lives, and monitor their cognitive health. It provides administrators with tools to "
    "identify users who may need support without compromising privacy.\n\n"
    "This report visually demonstrates the core features of the platform, outlining both the User Experience "
    "and the Administrative controls. The imagery clearly explains the platform's purpose and functionality."
)
doc.add_page_break()

doc.add_heading('2. The User Journey', level=1)
# 1
add_section(
    "Authentication", 
    "Users enter the platform through a secure login and registration portal.",
    "Figure 1: Landing and Authentication Page"
)
# 2
add_section(
    "The Dashboard", 
    "The central hub where users track their daily streak and follow the step-by-step 'Daily Journey'.",
    "Figure 2: Main User Dashboard"
)
# 3
add_section(
    "Daily Check-In", 
    "Users log their current emotional state using intuitive sliders and selectors.",
    "Figure 3: Daily Emotion Check-In"
)
# 4
add_section(
    "Reflection Journal", 
    "A private space for users to write down their thoughts and reflect on their day.",
    "Figure 4: Reflection Journal Interface"
)
# 5
add_section(
    "Scenario Assessment", 
    "Users are presented with real-life scenarios to test their decision-making and conflict resolution skills.",
    "Figure 5: Scenario Assessment Question"
)
# 6
add_section(
    "Scenario Feedback", 
    "Immediate, constructive feedback is provided based on the user's choice in the scenario.",
    "Figure 6: Scenario Assessment Feedback"
)
# 7
add_section(
    "Visual Memory Challenge", 
    "A cognitive mini-game designed to track short-term memory patterns over time.",
    "Figure 7: Memory Challenge Mini-Game"
)
# 8
add_section(
    "Reaction Time Challenge", 
    "A visual tracking game that monitors the user's cognitive reaction speeds.",
    "Figure 8: Reaction Time Challenge"
)
# 9
add_section(
    "Click Accuracy Challenge", 
    "The final cognitive test, measuring precision and focus.",
    "Figure 9: Click Accuracy Challenge"
)
# 10
add_section(
    "Lighthouse AI Companion", 
    "An AI assistant that analyzes the user's personal journal and check-ins to provide personalized insights and productivity tips.",
    "Figure 10: AI Companion Chat Interface"
)
# 11
add_section(
    "User Analytics", 
    "Users can view their historical data, tracking how their mood and cognitive performance change over time.",
    "Figure 11: Personal Progress and Analytics"
)

doc.add_page_break()
doc.add_heading('3. The Administrator Flow', level=1)

# 12
add_section(
    "Platform Analytics", 
    "Admins have a high-level overview of platform health, active users, and global completion rates.",
    "Figure 12: Admin Analytics Dashboard"
)
# 13
add_section(
    "User Management", 
    "A directory of all users, highlighting individuals flagged by the system based on negative check-in trends.",
    "Figure 13: Admin User Directory"
)
# 14
add_section(
    "The Care Panel", 
    "Admins can review attention signals, add users to a watchlist, and send gentle 'nudges', all without exposing private journal data.",
    "Figure 14: Administrative Care Panel"
)
# 15
add_section(
    "Scenario Management", 
    "A content management system allowing admins to create, edit, disable, or track the 50 built-in scenarios.",
    "Figure 15: Scenario Management Table"
)

doc.save('Lighthouse_Visual_Lab_Report.docx')
