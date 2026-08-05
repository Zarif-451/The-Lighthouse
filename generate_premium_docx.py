import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_bg(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

doc = docx.Document()
font_name = 'Segoe UI'

# Page Setup: Standard margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

normal_style = doc.styles['Normal']
normal_style.font.name = font_name
normal_style.font.size = Pt(10.5)
normal_style.font.color.rgb = RGBColor(31, 41, 55)

h1 = doc.styles['Heading 1']
h1.font.name = font_name
h1.font.size = Pt(15)
h1.font.bold = True
h1.font.color.rgb = RGBColor(30, 64, 175) # #1E40AF Deep Blue

h2 = doc.styles['Heading 2']
h2.font.name = font_name
h2.font.size = Pt(12)
h2.font.bold = True
h2.font.color.rgb = RGBColor(20, 184, 166) # #14B8A6 Teal

# --------------------------
# Clean Opening Header (Page 1 - Exact Version 16)
# --------------------------
title_p = doc.add_paragraph()
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after = Pt(2)
title_run = title_p.add_run('Project Lighthouse')
title_run.font.name = font_name
title_run.font.size = Pt(24)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(30, 64, 175)

sub_p = doc.add_paragraph()
sub_p.paragraph_format.space_before = Pt(0)
sub_p.paragraph_format.space_after = Pt(8)
sub_run = sub_p.add_run('Wellbeing and Self-Reflection Web Platform')
sub_run.font.name = font_name
sub_run.font.size = Pt(13)
sub_run.font.italic = True
sub_run.font.color.rgb = RGBColor(20, 184, 166)

p_line = doc.add_paragraph()
p_line.paragraph_format.space_before = Pt(2)
p_line.paragraph_format.space_after = Pt(14)
r_line = p_line.add_run('_________________________________________________________________________________')
r_line.font.name = font_name
r_line.font.color.rgb = RGBColor(203, 213, 225)

# --------------------------
# 1. Objective (Exact Version 16)
# --------------------------
doc.add_heading('1. Objective', level=1)
p_exec = doc.add_paragraph('The primary objective of Project Lighthouse is to develop a highly secure, private, and user-centric wellbeing and self-reflection platform. The application provides an elegant, responsive, and calm workspace where users can track key metrics, perform emotional checks, evaluate psychological scenarios, and record private thoughts without fearing data leakage. In parallel, the system provides an administrative layer that allows platform operators to monitor engagement trends and trigger care interventions for high-risk users, without ever exposing users\' personal journal texts.')
p_exec.paragraph_format.space_after = Pt(6)

doc.add_paragraph('Key objectives of the "Lighthouse" platform are:')

obj_bullets = [
    ('Structured Habit & Mood Tracking', 'Empowering individuals to monitor structured wellbeing metrics (sleep, mood, productivity, physical activity, water intake) on a daily basis.'),
    ('Guided Daily Journey (Core Loop)', 'Streamlining user engagement through a cohesive 4-step sequence (Check-in -> Scenario Assessment -> Visual Reflection -> Journal).'),
    ('Privacy-First Architecture', 'Ensuring complete data ownership. Row Level Security (RLS) restricts access to personal records. Under no circumstances can admin operators view private journal text.'),
    ('Platform-Side Attention Heuristics', 'Evaluating user check-in history to flag cases of consecutive low mood, sleep deprivation, or high social avoidance, allowing moderators to act proactively.'),
    ('Soft Platform Care', 'Enabling admins to log notes, initiate watchlist monitoring, and dispatch gentle \'nudge\' notifications to users needing attention.'),
    ('Real-time and Demo Analytics', 'Supplying interactive data visualizations powered by Chart.js and Recharts for both individual users and platform administrators.')
]

for title, desc in obj_bullets:
    bp = doc.add_paragraph(style='List Bullet')
    bp.paragraph_format.space_after = Pt(3)
    r_b = bp.add_run(f'{title}: ')
    r_b.font.name = font_name
    r_b.font.bold = True
    r_b.font.color.rgb = RGBColor(31, 41, 55)
    r_d = bp.add_run(desc)
    r_d.font.name = font_name
    r_d.font.color.rgb = RGBColor(55, 65, 81)

doc.add_paragraph()

# --------------------------
# 2. System Architecture & Sidebar Layout (Exact Version 16)
# --------------------------
doc.add_heading('2. System Architecture & Sidebar Layout', level=1)
doc.add_paragraph('The application is built around an integrated App Shell featuring a modern, collapsible Left Sidebar Navigation. This layout ensures fast, single-click access across all primary modules on desktop, tablet, and mobile displays without overcrowding the viewport header.')

# Page break so Section 3 starts cleanly on Page 2 (Exact Version 16)
doc.add_page_break()

# --------------------------
# 3. Detailed Module Walkthrough & Screenshots (Exact Version 16)
# --------------------------
doc.add_heading('3. Detailed Module Walkthrough & Screenshots', level=1)
doc.add_paragraph('This section presents a step-by-step walkthrough of every user interface in The Lighthouse. Each module is documented with its operational role, key feature set, and embedded visual screenshots.')

def add_module(num, title, tag, desc, fig, images=None):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Inches(4.5)
    tbl.columns[1].width = Inches(2.0)
    
    c1 = tbl.cell(0, 0)
    p = c1.paragraphs[0]
    p.paragraph_format.keep_with_next = True
    p_h2 = p.add_run(f'3.{num} {title}')
    p_h2.font.name = font_name
    p_h2.font.bold = True
    p_h2.font.size = Pt(11.5)
    p_h2.font.color.rgb = RGBColor(30, 64, 175)
    
    c2 = tbl.cell(0, 1)
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.keep_with_next = True
    r2 = p2.add_run(tag)
    r2.font.name = font_name
    r2.font.bold = True
    r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor(20, 184, 166)
    
    p_desc = doc.add_paragraph(desc)
    p_desc.runs[0].font.name = font_name
    p_desc.paragraph_format.space_after = Pt(4)
    p_desc.paragraph_format.keep_with_next = True
    
    if images:
        sub_labels = ['a', 'b', 'c', 'd']
        for idx, (img_path, sub_fig) in enumerate(images):
            img_p = doc.add_paragraph()
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_p.paragraph_format.space_before = Pt(4)
            img_p.paragraph_format.space_after = Pt(2)
            run_img = img_p.add_run()
            run_img.add_picture(img_path, width=Inches(5.8))
            
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_before = Pt(2)
            cap.paragraph_format.space_after = Pt(12 if idx < len(images)-1 else 16)
            sub_str = sub_labels[idx] if len(images) > 1 else ''
            r_cap = cap.add_run(f'Figure 3.{num}{sub_str}: {sub_fig}')
            r_cap.font.name = font_name
            r_cap.font.italic = True
            r_cap.font.size = Pt(9)
            r_cap.font.color.rgb = RGBColor(107, 114, 128)
    else:
        img_tbl = doc.add_table(rows=1, cols=1)
        img_tbl.style = 'Table Grid'
        tr = img_tbl.rows[0]._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '4320') # Exact Version 16 box height
        trHeight.set(qn('w:hRule'), 'exact')
        trPr.append(trHeight)
        c = img_tbl.cell(0, 0)
        cp = c.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.keep_with_next = True
        r_ph = cp.add_run('\n\n\n[PASTE SCREENSHOT HERE]')
        r_ph.font.name = font_name
        r_ph.font.size = Pt(9.5)
        r_ph.font.color.rgb = RGBColor(156, 163, 175)
        
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(4)
        cap.paragraph_format.space_after = Pt(16)
        r_cap = cap.add_run(f'Figure 3.{num}: {fig}')
        r_cap.font.name = font_name
        r_cap.font.italic = True
        r_cap.font.size = Pt(9)
        r_cap.font.color.rgb = RGBColor(107, 114, 128)

modules = [
    ('Authentication & User Onboarding', 'SECURITY & SESSION MANAGEMENT', 'The authentication module handles user sign-in and account creation securely through Supabase Auth. It enforces client-side validation, password metering, and securely manages JWT tokens to persist user sessions across browser reloads.', 'The Login and Sign-Up interface featuring form validation and secure user onboarding.'),
    ('Main Dashboard', 'CENTRAL CONTROL HUB', 'The Main Dashboard functions as the primary control center for the user. It aggregates streak counters, daily journey progress, recent entry statistics, and score indicators into a single unified view.', 'Upper Dashboard section containing streak statistics, score overview cards, and journey steps.'),
    ('Daily Check-In', 'EMOTIONAL ANALYTICS', 'The Mood Log allows users to quickly record their emotional state throughout the day using intuitive visual sliders. The accumulated data is automatically plotted into mood trend line charts.', 'Daily check-in widget with emotion selector sliders and visual distribution charts.'),
    ('Reflection Journal', 'REFLECTIVE WRITING', 'The Journal module enables users to log daily reflections in a private space. Upon entering text, the system parses linguistic patterns to compute real-time scores for emotional tone.', 'The Reflection Journal featuring interactive writing prompts and saved entries.'),
    ('Scenario Assessment', 'SITUATIONAL ANALYSIS', 'Users are presented with one of 50 curated real-life scenarios evaluating decision-making, conflict resolution, and coping mechanisms under pressure.', 'A multiple-choice Scenario Assessment questionnaire card.'),
    ('Scenario Feedback', 'COGNITIVE RESTRUCTURING', 'After completing a scenario assessment, immediate feedback is provided, highlighting alternative coping perspectives and constructive behavior recommendations.', 'Constructive feedback card displayed following scenario selection.'),
    ('Visual Memory Challenge', 'INTERACTIVE BRAIN TRAINING', 'A cognitive mini-game designed to track short-term memory patterns over time. Game performance metrics feed directly into the user\'s cognitive evaluation log.', 'The Visual Memory Challenge active game canvas area.'),
    ('Reaction Time Challenge', 'COGNITIVE RESPONSE', 'A visual tracking game that monitors the user\'s cognitive reaction speeds in milliseconds, measuring alertness, focus, and visual processing speed.', 'The Reaction Time Challenge interface with real-time response counters.'),
    ('Click Accuracy Challenge', 'PRECISION METRICS', 'The final cognitive test measures motor precision and sustained attention by tracking target hits and misses under time constraints.', 'The Click Accuracy Challenge target canvas area.'),
    ('Lighthouse AI Companion', 'VIRTUAL WELLNESS COMPANION', 'The AI Assistant provides a conversational chat interface powered by Groq LLM routing. It reads recent wellness context (moods, journal sentiment, game scores) to offer empathetic guidance.', 'Conversational chat window showing personalized reflection prompts and AI advice.'),
    ('User Analytics & History', 'AGGREGATED INSIGHTS', 'The Analytics page brings together all logged data sources—wellness index, average mood, cognitive scores—into a unified, side-by-side analytical dashboard.', 'Comprehensive Analytics dashboard displaying aggregated metric cards and trend charts.'),
    ('Admin Platform Analytics', 'ADMINISTRATIVE OVERSIGHT', 'A high-level dashboard exclusively for administrators. It aggregates platform health, active users, check-in completion rates, and system engagement.', 'Admin Analytics dashboard displaying global platform metrics.'),
    ('Admin User Directory', 'COMMUNITY HEALTH', 'A comprehensive directory of registered users. Administrators can search, filter by risk bands (Low, Moderate, High), and identify users needing support.', 'Admin User Directory with search and filtering capabilities.'),
    ('Admin Care Panel', 'INTERVENTION STRATEGIES', 'The Care Panel allows administrators to review attention signals, manage watchlists, and send gentle in-app nudges without ever exposing private journal text.', 'The Care Panel interface showing watchlist controls and nudge features.'),
    ('Scenario Management', 'CONTENT ADMINISTRATION', 'A full CRUD (Create, Read, Update, Delete) interface for the platform\'s scenario bank, allowing admins to edit, disable, or author scenarios.', 'Scenario Management table displaying all 50 platform scenarios with serial numbering.')
]

for i, m in enumerate(modules):
    if i == 0:
        sec1_images = [
            ('section3_1_landing.png', 'The Landing Page hero section with platform overview and navigation header.'),
            ('section3_1_signup.png', 'The User Registration and Account Creation form with interactive validation fields.')
        ]
        add_module(i+1, m[0], m[1], m[2], m[3], images=sec1_images)
    elif i == 1:
        sec2_images = [
            ('section3_2_dashboard.png', 'Upper Dashboard section containing streak statistics, score overview cards, and journey steps.')
        ]
        add_module(i+1, m[0], m[1], m[2], m[3], images=sec2_images)
    elif i == 2:
        sec3_images = [
            ('section3_3_checkin.png', 'Daily check-in widget with emotion selector sliders and visual distribution choices.')
        ]
        add_module(i+1, m[0], m[1], m[2], m[3], images=sec3_images)
    elif i == 3:
        sec4_images = [
            ('section3_4_journal.png', 'The Reflection Journal featuring interactive writing prompts and saved entries.')
        ]
        add_module(i+1, m[0], m[1], m[2], m[3], images=sec4_images)
    elif i == 4:
        sec5_images = [
            ('section3_5_scenario.png', 'A multiple-choice Scenario Assessment questionnaire card evaluating conflict resolution.')
        ]
        add_module(i+1, m[0], m[1], m[2], m[3], images=sec5_images)
    elif i == 5:
        sec6_images = [
            ('section3_6_feedback.png', 'Constructive feedback card displayed following scenario selection.')
        ]
        add_module(i+1, m[0], m[1], m[2], m[3], images=sec6_images)
    elif i == 6:
        sec7_images = [
            ('section3_7_memory.png', 'The Visual Memory Challenge active game canvas area.')
        ]
        add_module(i+1, m[0], m[1], m[2], m[3], images=sec7_images)
    elif i == 7:
        sec8_images = [
            ('section3_8_reaction_start.png', 'The Reaction Time Challenge intro card before initiating a test.'),
            ('section3_8_reaction_active.png', 'The active visual tracking canvas displaying the real-time response prompt.')
        ]
        add_module(i+1, m[0], m[1], m[2], m[3], images=sec8_images)
    elif i == 8:
        sec9_images = [
            ('section3_9_accuracy_start.png', 'The Click Accuracy Challenge intro card before starting the test.'),
            ('section3_9_accuracy_active.png', 'The Click Accuracy Challenge target canvas area with active precision target.')
        ]
        add_module(i+1, m[0], m[1], m[2], m[3], images=sec9_images)
    elif i == 9:
        sec10_images = [
            ('section3_10_companion.png', 'Conversational chat window showing personalized reflection prompts and AI advice.')
        ]
        add_module(i+1, m[0], m[1], m[2], m[3], images=sec10_images)
    elif i == 10:
        sec11_images = [
            ('section3_11_analytics.png', 'Comprehensive Analytics dashboard displaying aggregated metric cards and trend charts.')
        ]
        add_module(i+1, m[0], m[1], m[2], m[3], images=sec11_images)
    elif i == 11:
        sec12_images = [
            ('section3_12_admin_analytics.png', 'Admin Analytics dashboard displaying global platform metrics.')
        ]
        add_module(i+1, m[0], m[1], m[2], m[3], images=sec12_images)
    elif i == 12:
        sec13_images = [
            ('section3_13_user_directory.png', 'Admin User Directory with search and filtering capabilities.')
        ]
        add_module(i+1, m[0], m[1], m[2], m[3], images=sec13_images)
    elif i == 13:
        sec14_images = [
            ('section3_14_care_panel.png', 'The Care Panel interface showing watchlist controls and nudge features.')
        ]
        add_module(i+1, m[0], m[1], m[2], m[3], images=sec14_images)
    elif i == 14:
        sec15_images = [
            ('section3_15_scenarios.png', 'Scenario Management table displaying all 50 platform scenarios with serial numbering.')
        ]
        add_module(i+1, m[0], m[1], m[2], m[3], images=sec15_images)
    else:
        add_module(i+1, m[0], m[1], m[2], m[3])

# --------------------------
# 4. Technology Stack (Exact Version 16)
# --------------------------
doc.add_page_break()
doc.add_heading('4. Technology Stack & Implementation Details', level=1)
tech = doc.add_table(rows=1, cols=3)
tech.style = 'Table Grid'
hdr_tr = tech.rows[0]._tr.get_or_add_trPr()
hdr_tr.append(OxmlElement('w:cantSplit'))
hdr_cells = tech.rows[0].cells
hdr_cells[0].text = 'System Layer'
hdr_cells[1].text = 'Technology Selected'
hdr_cells[2].text = 'Implementation Role'

for c in hdr_cells:
    p = c.paragraphs[0]
    p.paragraph_format.keep_with_next = True
    p.runs[0].font.name = font_name
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    set_cell_bg(c, '1E40AF')

data = [
    ('Frontend Framework', 'Vanilla JavaScript, HTML5, CSS3', 'Renders a fast, lightweight single-page web application without heavy framework overhead.'),
    ('Styling & UI', 'Custom Vanilla CSS, Flexbox/Grid', 'Delivers modern dark-themed components, glassmorphism, and responsive layouts.'),
    ('Data Visualization', 'Chart.js & Recharts', 'Draws responsive area, line, bar, and radar charts for user and admin dashboards.'),
    ('Backend & Database', 'Supabase (PostgreSQL & Auth)', 'Manages user authentication, RLS security policies, and PostgreSQL database tables.'),
    ('AI Integration', 'Groq API (Llama 3.3 70B)', 'Provides high-speed conversational AI responses tailored to user reflection data.'),
    ('Export Engine', 'jsPDF & python-docx', 'Renders structured HTML and automated DOCX lab reports.')
]

for row_data in data:
    new_row = tech.add_row()
    row_tr = new_row._tr.get_or_add_trPr()
    row_tr.append(OxmlElement('w:cantSplit'))
    row_cells = new_row.cells
    for idx, text in enumerate(row_data):
        c = row_cells[idx]
        p = c.paragraphs[0]
        r = p.add_run(text)
        r.font.name = font_name
        r.font.size = Pt(9.5)
        if idx == 0:
            r.font.bold = True

doc.add_paragraph('\n')

# --------------------------
# 5. Application Flowchart (NOW DIRECTLY UNDERNEATH TECH STACK ON SAME PAGE!)
# --------------------------
# NO PAGE BREAK BEFORE SECTION 5 - Placed directly under Section 4 on the Tech Stack page!
p_fc_h1 = doc.add_heading('5. Application Flowchart', level=1)
p_fc_h1.paragraph_format.keep_with_next = True

p_fc_intro = doc.add_paragraph('The following flowchart illustrates the high-level architecture and user navigation paths throughout the Lighthouse platform, starting from authentication through to the core modules and administrative controls.')
p_fc_intro.paragraph_format.space_after = Pt(6)
p_fc_intro.paragraph_format.keep_with_next = True

p_img = doc.add_paragraph()
p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_img.paragraph_format.keep_with_next = True
run = p_img.add_run()
run.add_picture('flowchart_diagram.png', width=Inches(5.6)) # Sized to fit perfectly under the Tech Stack table!

cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap.paragraph_format.space_before = Pt(4)
cap.paragraph_format.space_after = Pt(12)
cap.paragraph_format.keep_with_next = True
r_cap = cap.add_run('Figure 5.1: High-Level Application Flowchart Diagram')
r_cap.font.name = font_name
r_cap.font.italic = True
r_cap.font.size = Pt(9)
r_cap.font.color.rgb = RGBColor(107, 114, 128)

# Page break so Section 6 ("Functionality") starts cleanly on the next page!
doc.add_page_break()

# --------------------------
# 6. Functionality (Exact Version 16)
# --------------------------
doc.add_heading('6. Functionality', level=1)

p_uf = doc.add_paragraph()
p_uf.paragraph_format.keep_with_next = True
r_uf = p_uf.add_run('User Panel Features')
r_uf.font.name = font_name
r_uf.font.bold = True
r_uf.font.size = Pt(12)
r_uf.font.color.rgb = RGBColor(20, 184, 166)

user_features_exact = [
    ('Daily Check-in', 'A form where users check in on structured indicators like hours of sleep, mood levels (scale of 1-5), energy levels, productivity, physical activity (None to High), water intake, and general notes. Data is stored in daily_checkins.'),
    ('Scenario Assessments (Slots 1-5)', 'Presents short, daily, randomized scenarios from the scenario bank. The user selects a coping option (A, B, C, or D), which maps to specific behavioral indicators. The results are recorded under scenario_responses. Up to 5 slots are allocated per day.'),
    ('Real-time Face Check', 'An ML-powered step that accesses the user\'s webcam (with full offline, skipped, or fallback support) to run clientside facial detection. It extracts expressions (happy, neutral, sad, surprised, angry, fearful, disgusted) and stores the dominant expression scores in behavioral_activity_results under \'face_check\'.'),
    ('Visual Reflections', 'Allows users to select an abstract photographic theme (e.g., Ocean, Forest, Storm Sky) that represents their current emotional context. Saves category choices under visual_reflections.'),
    ('Reflection Journal', 'A private, rich text field where users can write open reflections. Data is encrypted/secured in reflections, linked to the user\'s UUID, and completely shielded from administrative eyes.'),
    ('Insights Dashboard', 'Generates interactive trend charts using Chart.js for sleep patterns, mood shifts, productivity curves, selected visual themes, and facial expression distributions. Features a \'Demo Analytics\' mock toggle to show complete dashboards when history is thin.'),
    ('Lighthouse AI Companion', 'An intelligent, Groq LLM-powered assistant (Llama 3.3 70B) designed specifically for Lighthouse data interpretation. It securely reads the user\'s check-in trends, mood logs, reflection entries, and game scores to answer context-aware questions and provide personalized wellbeing insights.'),
    ('Personal Report Export', 'Integrates jsPDF on the client side, allowing the user to download a clean, formatted PDF report of their weekly or monthly check-in history and reflection frequency.')
]

for title, desc in user_features_exact:
    bp = doc.add_paragraph(style='List Bullet')
    bp.paragraph_format.space_after = Pt(3)
    r_b = bp.add_run(f'{title}: ')
    r_b.font.name = font_name
    r_b.font.bold = True
    r_b.font.color.rgb = RGBColor(31, 41, 55)
    r_d = bp.add_run(desc)
    r_d.font.name = font_name
    r_d.font.color.rgb = RGBColor(55, 65, 81)

doc.add_paragraph()

p_af = doc.add_paragraph()
p_af.paragraph_format.keep_with_next = True
r_af = p_af.add_run('Admin Panel Features')
r_af.font.name = font_name
r_af.font.bold = True
r_af.font.size = Pt(12)
r_af.font.color.rgb = RGBColor(20, 184, 166)

admin_features_exact = [
    ('Platform Health & Analytics', 'Displays overall usage counts, user registration numbers, total journal submissions, completed scenarios, and active participant numbers over time.'),
    ('Flagged User Care Panel', 'Aggregates attention flags without exposing private texts. Admins can place high-risk users on a 7-day watchlist, log private case notes, and trigger gentle soft nudges shown on user home screens.'),
    ('Care Flagging Logic', 'Heuristics evaluate recent data logs: (1) low mood streaks (>=3 checkins under \'Neutral\' in the last 5 checkins), (2) sleep deprivation (<5 hours in multiple logs), (3) social avoidance (selecting avoidant options in scenario responses), and (4) inactivity gaps (>=4 days without check-ins).'),
    ('Scenario Bank Manager', 'Provides interface elements to create new scenario items, edit existing stories, define choices (A, B, C, D), and toggle items as active or inactive.'),
    ('High-Value PDF Packs', 'Generates structured PDF packages on the fly via jsPDF, including Platform Snapshots, Engagement Logs, Flagged User Case Summaries, and User Cohort breakdowns.')
]

for title, desc in admin_features_exact:
    bp = doc.add_paragraph(style='List Bullet')
    bp.paragraph_format.space_after = Pt(3)
    r_b = bp.add_run(f'{title}: ')
    r_b.font.name = font_name
    r_b.font.bold = True
    r_b.font.color.rgb = RGBColor(31, 41, 55)
    r_d = bp.add_run(desc)
    r_d.font.name = font_name
    r_d.font.color.rgb = RGBColor(55, 65, 81)

doc.add_paragraph('\n')

# --------------------------
# 7. Conclusion (Exact Version 16)
# --------------------------
doc.add_heading('7. Conclusion', level=1)
doc.add_paragraph('The Lighthouse platform successfully integrates cognitive tracking, interactive brain training, emotional reflection, and conversational AI into a unified web application. The clean modular architecture, responsive layout, elegant font choices, and robust administrative tools make it a highly effective and visually appealing system for personal wellbeing management.')

import time
filename = 'Lighthouse_Premium_Lab_Report_v22.docx'
saved = False
for attempt in range(3):
    try:
        doc.save(filename)
        print(f"Successfully generated {filename}")
        saved = True
        break
    except PermissionError:
        time.sleep(1)

if not saved:
    raise PermissionError(f"File {filename} is open in Microsoft Word! Please close Word so I can update v22.")
